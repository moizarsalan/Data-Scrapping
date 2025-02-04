import requests
from bs4 import BeautifulSoup
import csv
import os
from docx import Document

# Base URL of the website
base_url = "http://quotes.toscrape.com/page/{}/"

# Function to get the next available file number
def get_next_file_number(extension):
    i = 1
    while os.path.exists(f"quotes_{i}.{extension}"):
        i += 1
    return i

# Get the next available file names
csv_file_number = get_next_file_number("csv")
docx_file_number = get_next_file_number("docx")

csv_filename = f"quotes_{csv_file_number}.csv"
docx_filename = f"quotes_{docx_file_number}.docx"

# Create a new Word document
doc = Document()
doc.add_heading('Scraped Quotes', level=1)

# Open a CSV file to write the data
with open(csv_filename, mode='w', newline='', encoding='utf-8', errors='ignore') as file:
    writer = csv.writer(file, quoting=csv.QUOTE_ALL)
    writer.writerow(['Quote', 'Author', 'Tags'])  # Write the header row

    # Keep track of the number of quotes scraped
    quote_count = 0
    page_number = 1
    quote_index = 1  # Numbering for the quotes

    while quote_count < 500:
        # Construct the URL for the current page
        url = base_url.format(page_number)
        response = requests.get(url)
        soup = BeautifulSoup(response.text, 'html.parser')

        # Find all quotes on the page
        quotes = soup.find_all('div', class_='quote')

        # If no quotes are found, stop the loop (end of the pages)
        if not quotes:
            break

        # Scrape each quote on the page
        for quote in quotes:
            text = quote.find('span', class_='text').get_text(strip=True)  # Extract the quote text
            author = quote.find('small', class_='author').get_text(strip=True)  # Extract the author name
            tags = [tag.get_text(strip=True) for tag in quote.find_all('a', class_='tag')]  # Extract the tags

            # Write the data into the CSV file
            writer.writerow([text, author, ', '.join(tags)])

            # Add the numbered quote to the Word document
            doc.add_paragraph(f'{quote_index}. "{text}"', style='Quote')
            doc.add_paragraph(f'— {author}', style='Normal')
            doc.add_paragraph(f'Tags: {", ".join(tags)}\n', style='Normal')

            # Increase the quote count
            quote_count += 1
            quote_index += 1  # Increment numbering

            # Stop if we have scraped 500 quotes
            if quote_count >= 500:
                break

        # Move to the next page
        page_number += 1

# Save the Word document
doc.save(docx_filename)

print(f"500 quotes have been scraped and saved to '{csv_filename}' and '{docx_filename}'.")
